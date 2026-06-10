#!/usr/bin/env python3
"""PA02 6-page report generator — PA02_최종보고서_안승현.docx"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
TEMPLATE = HERE / "PA02_Optimization_Report_Template.docx"
FIG = HERE / "figures"
OUT = HERE / "PA02_최종보고서_안승현.docx"
KF = "맑은 고딕"

# ─── helpers ─────────────────────────────────────────────────────────────────

def _kfont(run, name=KF, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rf.set(qn(attr), name)
    if size:  run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color: run.font.color.rgb = color

def _pf(p, sb=2, sa=2, li=None, al=None):
    if al is not None: p.alignment = al
    f = p.paragraph_format
    f.space_before = Pt(sb); f.space_after = Pt(sa)
    if li: f.left_indent = Inches(li)

def clear_body(doc):
    body = doc.element.body
    for c in list(body):
        if c.tag != qn("w:sectPr"): body.remove(c)

def P(doc, text, sz=10, bold=False, italic=False, al=None,
      sb=2, sa=2, li=None, color=None):
    p = doc.add_paragraph()
    _pf(p, sb, sa, li, al)
    if text:
        r = p.add_run(text)
        _kfont(r, size=sz, bold=bold, italic=italic, color=color)
    return p

def mixed(doc, segs, sz=10, al=None, sb=2, sa=2, li=None):
    """segs: [(text, bold, color?), ...]"""
    p = doc.add_paragraph()
    _pf(p, sb, sa, li, al)
    for seg in segs:
        txt, bld = seg[0], seg[1]
        col = seg[2] if len(seg) > 2 else None
        r = p.add_run(txt)
        _kfont(r, size=sz, bold=bld, color=col)
    return p

def H(doc, text, lv=1):
    sz = {1: 14, 2: 11.5, 3: 10.5}[lv]
    sb = {1: 9, 2: 6, 3: 4}[lv]
    p = doc.add_paragraph(style="Heading %d" % lv)
    _pf(p, sb=sb, sa=3)
    r = p.add_run(text)
    _kfont(r, size=sz, bold=True)
    return p

def bullet(doc, label, body, sz=10):
    p = doc.add_paragraph()
    _pf(p, sb=1, sa=1, li=0.18)
    r = p.add_run("• "); _kfont(r, size=sz, bold=True)
    if label:
        r2 = p.add_run(label + " "); _kfont(r2, size=sz, bold=True)
    r3 = p.add_run(body); _kfont(r3, size=sz)
    return p

def shade(cell, hexcol):
    tcp = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hexcol); tcp.append(s)

def tbl(doc, hdrs, rows, wids=None, sz=9, hfill="1F3864"):
    t = doc.add_table(rows=1, cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hc = t.rows[0].cells
    for i, h in enumerate(hdrs):
        hc[i].text = ""
        p = hc[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _pf(p, sb=1, sa=1)
        r = p.add_run(h); _kfont(r, size=sz, bold=True,
                                  color=RGBColor(0xFF,0xFF,0xFF))
        shade(hc[i], hfill)
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            p = cs[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _pf(p, sb=1, sa=1)
            bld = v.startswith("**")
            txt = v[2:] if bld else v
            r = p.add_run(txt); _kfont(r, size=sz, bold=bld)
    if wids:
        for i,w in enumerate(wids):
            for row in t.rows: row.cells[i].width = Inches(w)
    return t

def fig(doc, fname, cap, width=5.1):
    p = doc.add_paragraph()
    _pf(p, sb=4, sa=1, al=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(FIG/fname), width=Inches(width))
    c = doc.add_paragraph()
    _pf(c, sb=0, sa=6, al=WD_ALIGN_PARAGRAPH.CENTER)
    r = c.add_run(cap)
    _kfont(r, size=8.5, bold=True, color=RGBColor(0x40,0x40,0x40))

DGRAY = RGBColor(0x40,0x40,0x40)
DBLUE = RGBColor(0x1A,0x56,0x96)
DRED  = RGBColor(0xA8,0x1E,0x2C)

# ─── build document ──────────────────────────────────────────────────────────
doc = Document(str(TEMPLATE))
clear_body(doc)

# ══════════════════════════════════════════
# 표지 (Title block)
# ══════════════════════════════════════════
P(doc, "실시간 데이터 해석 — PA02 최종 보고서",
  sz=21, bold=True, al=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=3)
P(doc, "Fast Correlative Scan Matcher 모듈 아키텍처 관점의 오케스트레이션 계층 병렬 고속화",
  sz=11, al=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=10, color=DGRAY)

meta = [
    ("담당교수", "스마트모빌리티공학과 심인욱 교수"),
    ("학번 / 이름", "12223637 / 안승현"),
    ("제출일", "2026-06-10"),
]
mt = doc.add_table(rows=len(meta), cols=2)
mt.style = "Table Grid"
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate(meta):
    mt.rows[i].cells[0].text=""
    mt.rows[i].cells[1].text=""
    rk = mt.rows[i].cells[0].paragraphs[0].add_run(k)
    rv = mt.rows[i].cells[1].paragraphs[0].add_run(v)
    _kfont(rk, size=9.5, bold=True)
    _kfont(rv, size=9.5)
    mt.rows[i].cells[0].width = Inches(1.5)
    mt.rows[i].cells[1].width = Inches(4.0)
    for c in mt.rows[i].cells:
        c.paragraphs[0].paragraph_format.space_before=Pt(2)
        c.paragraphs[0].paragraph_format.space_after=Pt(2)
    shade(mt.rows[i].cells[0], "E8EDF5")

P(doc, "", sb=4, sa=0)   # spacer

# ══════════════════════════════════════════
# Abstract
# ══════════════════════════════════════════
H(doc, "Abstract", lv=2)
P(doc,
  "본 보고서는 Google Cartographer 기반 2D Fast Correlative Scan Matcher의 오케스트레이션 계층 "
  "(make_cand, Score, Branch)을 Jetson Nano 실환경에서 최적화한 PA02 과제의 결과를 정리한다. "
  "PA01에서 score_all 커널을 L9 GPU 하이브리드(threshold=256)로 고정, 61% 가속을 달성한 뒤 "
  "bag 프로파일링을 통해 matcher_scope에 잔여 병목 56,667 ms(62.6%)가 존재함을 실측하였다. "
  "Phase 1(make_cand reserve) → Phase 2(Branch buffer reuse) → Phase 3(Score scan-bucket) 순으로 "
  "누적 구조 개선을 적용하여 모듈 KPI [match] cumulative를 90,482 ms → 87,866 ms(−2.9%)로 단축하였으며, "
  "best_score=0.783을 전 단계에서 유지해 정확도 회귀가 없음을 검증하였다. "
  "본 연구의 방법론적 핵심은 격리 마이크로벤치가 아닌 roslaunch 기반 bag 구동 환경의 누적 시간을 "
  "유일한 KPI로 삼은 것이며, 이를 통해 실제 SLAM 파이프라인에서 GPU crossover가 n=256 수준으로 "
  "내려온다는 현상을 독자적으로 규명하고 검증하였다.",
  sz=10, sb=0, sa=6)

# ══════════════════════════════════════════
# 제1장 서론
# ══════════════════════════════════════════
H(doc, "제 1 장.  서론 및 문제 정의")
H(doc, "1.1  PA01 이후 남은 과제와 최적화 대상 선정", lv=2)
P(doc,
  "PA01을 통해 단일 커널 score_all()은 성공적으로 고속화하였으나, 상위 알고리즘인 "
  "Branch-and-Bound(B&B) 트리의 호출 구조, 동적 메모리 할당, 정렬 오버헤드 등 "
  "오케스트레이션 계층의 병목이 그대로 잔존하였다. 추측에 의한 코드 수정을 배제하기 위해 "
  "pa02_timing.h에 chrono 기반 타이머를 삽입하여 L0 Baseline의 태그별 누적 시간을 실측하였다.",
  sz=10, sb=0, sa=3)
mixed(doc, [
    ("실측 결과 전체 [match] 시간 중 ", False),
    ("matcher_scope = match_total − score_all_total", True),
    ("로 정의되는 오케스트레이션 영역이 56,667 ms(62.6%)로 지배적 병목임을 확인하고, "
     "Branch(60.5%), Score(66.2%), make_cand(1.3%)를 PA02의 3대 최적화 타깃으로 확정하였다.", False),
], sz=10, sb=0, sa=3)
P(doc,
  "아래 호출 트리는 최적화 범위를 명시한다. score_all(PA01 고정)은 Score 내부에서만 호출되며, "
  "PA02는 그 상위의 make_cand, Score, Branch 세 함수만을 대상으로 한다.",
  sz=10, sb=0, sa=2)
doc.add_paragraph()  # small space before code
_p = doc.paragraphs[-1]; _p.paragraph_format.space_after=Pt(0)
for line in [
    "MatchWithWindow()          ← [match] KPI",
    "  ├─ MakeLowCands → make_cand()    [PA02 Phase 1]",
    "  ├─ Score()      → score_all()    [PA01 고정]",
    "  └─ Branch()     → Score() 재귀   [PA02 Phase 2/3]",
]:
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(0)
    cp.paragraph_format.left_indent  = Inches(0.3)
    ppr = cp._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),"F2F2F2")
    ppr.append(shd)
    r = cp.add_run(line)
    r.font.name = "Consolas"; r.font.size = Pt(8.5)
    rpr2 = r._element.get_or_add_rPr()
    rf2 = OxmlElement("w:rFonts")
    rf2.set(qn("w:ascii"),"Consolas"); rf2.set(qn("w:hAnsi"),"Consolas")
    rpr2.append(rf2)
P(doc, "", sb=3, sa=0)

H(doc, "1.2  KPI 및 정확도 지표 정의", lv=2)
P(doc,
  "단순 벽시계 시간이 아닌 ROS 네트워크·launch 지연을 배제한 순수 Matcher 연산 누적인 "
  "[match] cumulative (ms)를 최종 KPI로 삼았다. 최적화 전 과정에서 B&B가 찾은 최고 정합 점수 "
  "best_score(기준값 0.783)와 coarse_n(기준값 3840)이 Baseline과 완벽히 일치하는지 매 Phase마다 "
  "교차 검증하였다. 이는 전역 SLAM RMSE가 아닌 '동일 bag·동일 map에서 Matcher 출력이 "
  "Baseline과 동일한가'를 보는 Regression sanity check이다.",
  sz=10, sb=0, sa=4)

# ══════════════════════════════════════════
# 제2장 환경·워크플로·측정 철학
# ══════════════════════════════════════════
H(doc, "제 2 장.  실험 환경, 개발 워크플로 및 측정 철학")
H(doc, "2.1  실험 환경", lv=2)
tbl(doc,
    ["항목", "사양"],
    [("SoC", "Jetson Nano — 4× Cortex-A57, Maxwell 128-core GPU, 4 GB LPDDR4 (CPU·GPU 대역폭 공유)"),
     ("소프트웨어", "CUDA 10.2 / ROS Melodic / Docker student_19 / catkin_ws"),
     ("ROS Master", "192.168.0.106:11311 / ROS_IP 192.168.0.104"),
     ("맵·Bag", "467×314 cells (0.05 m/cell, ~147 KB) / scan.bag (LaserScan, p≈1081)"),
     ("PA01 고정", "PA01_OPT_LEVEL=9, PA01_USE_GPU=ON, PA01_GPU_THRESHOLD=256"),
     ("PA02 프로파일", "PA02_PROFILE=ON, pa02_timing.h chrono 태그"),],
    wids=[1.4, 5.1], sz=9)
P(doc,"",sb=3,sa=0)

H(doc, "2.2  Git 기반 개발 파이프라인", lv=2)
P(doc,
  "PA02는 fast_matcher.cpp 전체 구조를 반복 수정하는 과제이므로 SSH 이중 접속·nano 편집 방식으로는 "
  "빠른 실험(Iteration)이 불가능하였다. ~/.ssh/config에 ProxyJump를 설정하여 로컬에서 "
  "ssh jetson-nano-19 명령 한 번으로 Docker 컨테이너 터미널에 직결되도록 하였다. "
  "로컬 IDE에서 전체 프로젝트 트리를 조망하며 코드를 수정한 뒤 GitHub push → Jetson git pull → "
  "catkin_make → pa02_bag_profile.sh 로 이어지는 루프를 자동화하였다. "
  "이 워크플로는 upstream ref/cartographer 참고 실험까지 모든 구조 변경을 짧은 주기로 "
  "재현 가능하게 유지하는 기반이 되었다.",
  sz=10, sb=0, sa=3)

H(doc, "2.3  측정 철학 — bag 누적 시간이 정확한 KPI인 이유", lv=2)
mixed(doc,[
    ("본 연구의 가장 핵심적인 방법론적 선택은 ", False),
    ("격리 마이크로벤치가 아닌 roslaunch 기반 bag replay의 태그별 누적 시간", True),
    ("을 유일한 KPI로 삼은 것이다. 이 선택의 타당성은 세 가지 실측 반례로 뒷받침된다.", False),
    ], sz=10, sb=0, sa=2)

bullet(doc,"[반례 1] 호출 수 피드백:",
       "score_all이 빨라지면 Branch/Score가 더 깊이 탐색해 호출 수가 늘어난다. "
       "PA02 L0→L3에서 score_all calls +4.6% 증가가 이를 보여준다. "
       "함수 단위 avg만으로는 이 피드백 효과를 발견할 수 없다.")
bullet(doc,"[반례 2] crossover 역전:",
       "연속 마이크로벤치에서 CPU(OpenMP) crossover는 n≥2048이었으나, "
       "실제 bag에서 n=256 CPU ~2.00 ms vs GPU ~0.86 ms로 GPU가 2.3× 우세하였다. "
       "마이크로벤치 단독 결론(GPU 불필요)은 bag에서 score_all cumulative +38% 악화를 초래한다.")
bullet(doc,"[반례 3] 정확도 미검출:",
       "ShrinkToFit 실험에서 마이크로벤치 속도는 개선되었으나 "
       "bag에서만 best_score 0.783→0.748 정확도 회귀가 검출되었다.")

P(doc,
  "따라서 본 보고서는 마이크로벤치를 가설 선별 도구로만 활용하고, "
  "모든 최종 채택·기각 판단은 bag cumulative 실측 데이터로 결정하였다.",
  sz=10, sb=3, sa=3)

fig(doc, "fig04_microbench_vs_bag.png",
    "Fig. 1  마이크로벤치(좌) vs 실제 bag 환경(우) — Phase 3에서 승자가 역전된다 "
    "(cpu_score: μbench 1위 → bag 4위, +4.4%)",
    width=5.3)

# ══════════════════════════════════════════
# 제3장 Phase 0 병목 분해
# ══════════════════════════════════════════
H(doc, "제 3 장.  Phase 0 — 정량적 병목 분해 및 최적화 로드맵")
H(doc, "3.1  모듈 KPI 분해 및 최적화 상한선(Ceiling) 설정", lv=2)
P(doc,
  "동일 scan.bag + PA01 L9 빌드로 L0 Baseline bag를 구동하고 pa02_analyze_profile.py로 "
  "태그별 누적 시간을 자동 분해하였다. 결과 [match] 90,482 ms 중 PA01 score_all 커널은 "
  "33,814 ms(37.4%)이며, PA02 최적화 상한선은 matcher_scope 56,667 ms(62.6%)임을 확정하였다.",
  sz=10, sb=0, sa=3)

fig(doc, "fig01_l0_match_decomposition.png",
    "Fig. 2  L0 bag [match] 90,482 ms 시간 분해 — matcher_scope 56,667 ms(62.6%)가 PA02 타깃",
    width=5.0)

H(doc, "3.2  함수별 런타임 및 Stratum 분석", lv=2)
tbl(doc,
    ["함수", "L0 누적 (ms)", "[match] 대비", "주요 병목 원인", "최적화 Phase"],
    [("Branch()", "54,743", "60.5%", "재귀마다 child vector 힙 할당, empty quadrant 미스킵, depth=3 집중(34,679 ms)", "Phase 2"),
     ("Score()", "59,869", "66.2%", "스캔마다 cand 전체 선형 스캔, 매 호출 vector 재할당, Score_orch 25,955 ms", "Phase 3"),
     ("make_cand()", "1,216", "1.3%", "31,755회 호출 중 push_back 재할당 스트레스 누적", "Phase 1"),
     ("score_all()", "33,814", "37.4%", "PA01 L9 GPU hybrid 완료 — PA02 변경 없음", "고정"),],
    wids=[1.2, 1.15, 0.85, 2.6, 0.85], sz=8.5)
P(doc,"",sb=3,sa=0)
P(doc,
  "중첩 타이머 주의: Branch ⊃ Score ⊃ score_all 구조이므로 누적 합산으로 단순 비교하면 "
  "이중 계상이 발생한다. 본 보고서는 matcher_scope = match − score_all, "
  "Score_orchestration = Score − score_all 분리식을 사용하였다.",
  sz=9.5, sb=2, sa=3, color=DGRAY)

H(doc, "3.3  최적화 3단계 로드맵 확정", lv=2)
P(doc,
  "리스크와 파급력을 고려하여 Phase 순서를 아래와 같이 확정하였다. "
  "Phase 1(make_cand)은 절대 비중은 작으나 side-effect가 없어 "
  "실험 파이프라인 검증 및 regression 기준선 확립에 최적이다. "
  "Phase 2(Branch)는 depth=3 stratum이 35 s로 B&B 구조 자체를 개선한다. "
  "Phase 3(Score)는 오케스트레이션 26 s를 직접 겨냥하므로 최대 KPI 기여가 기대된다.",
  sz=10, sb=0, sa=4)

# ══════════════════════════════════════════
# 제4장 Phase 1-2
# ══════════════════════════════════════════
H(doc, "제 4 장.  Phase 1–2 — CPU 오케스트레이션 최적화 구현")
H(doc, "4.1  Phase 1: make_cand 격자 생성 최적화 (L1)", lv=2)
P(doc,
  "탐색 창 크기 (max_x − min_x) / step + 1 로 최종 후보 수 n_out을 선계산할 수 있으므로, "
  "루프 진입 전 cx->reserve(cx->size() + n_out)를 수행하여 "
  "push_back에 의한 vector growth·realloc 오버헤드를 원천 차단하였다.",
  sz=10, sb=0, sa=2)
mixed(doc,[
    ("OpenMP 임계값 전략:  ", True),
    ("bag Hot Path는 후보 수 256(16×16)이다. 이 구간에서 OpenMP를 발동하면 "
     "스레드 fork-join 오버헤드로 0.0020 ms → 0.0042 ms(×2.1 악화)가 측정되었다. "
     "PA02_MAKE_CAND_OMP_MIN=512를 채택하여 hot path에서 OpenMP가 발동되지 않도록 제어하였다.",
     False)], sz=10, sb=0, sa=3)

H(doc, "4.2  Phase 2: Branch 재귀 및 동적 할당 최적화 (L2)", lv=2)
bullet(doc,"구현 A — thread_local buffer reuse:",
       "B&B 재귀마다 발생하는 child vector 힙 할당을 제거하기 위해 "
       "static thread_local std::vector<Cand> child를 선언하고 clear()+reserve(4)로 재사용하였다.")
bullet(doc,"구현 B — empty quadrant skip:",
       "bounds 밖 사분면(dx/dy가 max_x/max_y를 초과)은 child 생성을 건너뛰어 "
       "depth=3 계층의 무효 탐색을 줄였다.")
mixed(doc,[
    ("Sibling OMP 기각:  ", True),
    ("CPU-only microbench에서 Branch 하위 루프에 OpenMP를 추가하면 "
     "28.4 ms → 22.2 ms(−22%) 개선이 확인되었다. 그러나 GPU 빌드 환경에서는 "
     "다중 OMP 스레드가 전역 CUDA singleton(DeviceBuffers)에 동시 접근하여 "
     "d_cx/d_score 덮어쓰기·stream interleave로 인한 data race가 발생한다. "
     "PA02_BRANCH_OMP_MIN=999999 + #ifndef PA01_USE_GPU compile guard로 "
     "GPU 빌드에서 compile-time 차단하였다.",
     False)], sz=10, sb=2, sa=3)

P(doc,
  "Phase 1–2 누적 결과: [match] cumulative L0 90,482 ms → L2 89,926 ms(−0.6%). "
  "best_score=0.783, coarse_n=3840 완벽 유지. score_all regression 없음(±2% 이내).",
  sz=10, sb=0, sa=4, color=DBLUE)

# ══════════════════════════════════════════
# 제5장 Phase 3 + hybrid
# ══════════════════════════════════════════
H(doc, "제 5 장.  Phase 3 — Score 파이프라인 개조 및 Hybrid 아키텍처 확정")
H(doc, "5.1  Phase 3: Score 1-pass scan-bucket 파이프라인 (L3)", lv=2)
P(doc,
  "L0 Score()는 스캔마다 cand 전체(O(n_cand × n_scans))를 선형 스캔하며 매번 ids/cx/cy vector를 "
  "힙에 새로 할당하여 score_all을 호출하였다. Score_orchestration이 25,955 ms(Score의 43%)를 "
  "점유한 원인이다.",
  sz=10, sb=0, sa=2)
P(doc,
  "L3에서는 candidates를 scan index 기준으로 1-pass bucket 분류(O(n_cand)) 후 "
  "static thread_local cx/cy/scores 버퍼를 재사용하여 scan별 score_all을 호출하였다. "
  "empty scan skip으로 불필요한 score_all 호출도 제거하였다.",
  sz=10, sb=0, sa=3)

H(doc, "5.2  CPU vs GPU 역할 분해 확인", lv=2)
P(doc,
  "PA02 L3는 CPU 오케스트레이션만 변경하며 score_all(PA01) 커널을 건드리지 않는다. "
  "bag score_all path 분해 결과 cuda 경로 ~81%, n4 ~18%, interchange <1%이며, "
  "GPU는 PA01 hybrid가 이미 담당하고 있다. PA02 별도 GPU 레이어 추가는 "
  "이미 GPU가 처리하는 coarse kernel 위에 중복 레이어를 얹는 것이므로 ROI가 낮다고 판단하였다.",
  sz=10, sb=0, sa=3)

H(doc, "5.3  Phase 3 hybrid 결정 실험 및 마이크로벤치 왜곡 분석", lv=2)
P(doc,
  "PA02_OPT_LEVEL과 PA01_GPU_THRESHOLD 조합을 바꾸며 동일 bag를 4회 실측하였다.",
  sz=10, sb=0, sa=2)
tbl(doc,
    ["Variant", "PA02", "GPU Threshold", "[match] (ms)", "Δ vs hybrid", "best_score"],
    [("**hybrid_prod (채택)", "L3", "256", "**87,866", "기준", "0.783"),
     ("gpu_aggr", "L3", "64",  "87,713", "−0.2%", "0.783"),
     ("legacy_gpu", "L2", "256","89,926", "+2.3%", "0.783"),
     ("cpu_score (GPU OFF)", "L3", "999999","91,750", "+4.4%", "0.783"),],
    wids=[1.45, 0.55, 0.9, 0.95, 0.8, 0.75], sz=8.5)
P(doc,"",sb=2,sa=0)
P(doc,
  "cpu_score는 격리 마이크로벤치에서 36.6 ms로 가장 빠르게 측정되었으나, "
  "bag에서는 91,750 ms(+4.4%)로 꼴찌였다. GPU가 꺼지면 coarse(n=256)가 OpenMP로 처리되어 "
  "score_all cumulative가 34,786 ms → 47,325 ms(+36%)로 급증하기 때문이다. "
  "마이크로벤치 승자를 채택했을 경우 발생하는 KPI 손실의 정량적 근거이다.",
  sz=10, sb=2, sa=4)

fig(doc, "fig06_phase_progression.png",
    "Fig. 3  PA02 Phase 진행에 따른 [match] 누적 시간 감소 (L0→L3, best_score=0.783 전 단계 유지)",
    width=5.1)

# ══════════════════════════════════════════
# 제6장 upstream 실험
# ══════════════════════════════════════════
H(doc, "제 6 장.  Upstream Cartographer 구조 이식 실험 및 기각 분석")
H(doc, "6.1  실험 설계", lv=2)
P(doc,
  "ref/cartographer/fast_correlative_scan_matcher_2d.cc를 read-only diff하여 "
  "ShrinkToFit(E1), exact reserve(E2), sort-skip(E3) 세 구조를 추출하고 "
  "PA02_OPT_LEVEL=4~5에 단계별로 이식하였다. pa02_structural_sweep.sh로 "
  "동일 bag를 자동 반복 측정하여 KPI와 정확도를 동시에 추적하였다.",
  sz=10, sb=0, sa=3)

H(doc, "6.2  E1: ShrinkToFit — 정확도 회귀로 기각", lv=2)
P(doc,
  "upstream은 discrete scan의 in-bounds와 linear window 교집합으로 탐색창을 축소한다. "
  "동일 로직을 MakeBounds local path에 적용하자 다음 결과가 실측되었다.",
  sz=10, sb=0, sa=2)
tbl(doc,
    ["Level", "[match] (ms)", "coarse_n", "best_score", "판정"],
    [("L3 (baseline)", "88,276", "3840", "**0.783", "기준"),
     ("L4 + ShrinkToFit", "91,326", "**3600", "**0.748", "⚠ 정확도 회귀"),
     ("L5", "91,853", "3600", "0.748", "회귀 지속"),
     ("L6", "92,567", "3600", "0.777", "회귀 지속"),],
    wids=[1.1, 1.15, 0.85, 0.85, 1.2], sz=9)
P(doc,"",sb=2,sa=0)
P(doc,
  "원인: 본 패키지 score_all은 OOB(Out-of-Bounds) 포인트를 0으로 처리하여 "
  "부분 중첩 후보도 유효 pose로 허용하는 semantics를 가진다. "
  "ShrinkToFit은 이 부분 중첩 영역을 강제 제거하므로 최적 pose 후보가 탐색 공간 밖으로 잘린다. "
  "또한 마이크로벤치에서는 L6(34.2 ms)이 L3(48.0 ms)보다 빠르게 나타났으나 "
  "bag KPI·정확도 모두 악화 — microbench ≠ bag 패턴의 세 번째 사례이다. "
  "정확도를 희생한 속도 향상은 공학적 가치가 없다고 판단, 코드를 완전 롤백하였다.",
  sz=10, sb=2, sa=3)

fig(doc, "fig05_shrinktofit_failure.png",
    "Fig. 4  ShrinkToFit 이식 실험 — match 악화 + best_score 0.783→0.748 정확도 회귀 실측",
    width=5.1)

H(doc, "6.3  E2/E3: exact reserve·sort-skip — 미채택", lv=2)
P(doc,
  "ShrinkToFit revert 후 L4(exact reserve), L5(sort-skip + presize)를 별도 재측정하였다. "
  "마이크로벤치에서는 −6.4% 개선이 있었으나 bag KPI는 ±0.6% 노이즈 범위(L3 대비 +0.4~0.6%)로 "
  "유의미한 개선이 없었다. 코드 복잡도 증가 대비 이득이 없어 L3(Phase 1~3)를 최종 유지하였다.",
  sz=10, sb=0, sa=4)

# ══════════════════════════════════════════
# 제7장 종합 결론
# ══════════════════════════════════════════
H(doc, "제 7 장.  종합 성능 비교 및 결론")
H(doc, "7.1  PA02 단계별 누적 결과 요약", lv=2)
tbl(doc,
    ["Level", "핵심 기법", "[match] (ms)", "Δ vs L0", "Score_orch (ms)", "best_score"],
    [("L0 (Baseline)", "PA01 L9 GPU Hybrid 고정", "90,482", "기준", "25,955", "0.783"),
     ("L1", "make_cand reserve + OMP 제어", "90,701", "+0.2% (노이즈)", "—", "0.783"),
     ("L2", "Branch thread_local buffer reuse", "89,926", "−0.6%", "—", "0.783"),
     ("**L3 (최종)", "Score 1-pass scan-bucket", "**87,866", "**−2.9%", "**22,235", "**0.783"),],
    wids=[0.85, 2.0, 0.9, 1.05, 1.0, 0.75], sz=8.5)
P(doc,"",sb=3,sa=0)

H(doc, "7.2  PA01 + PA02 end-to-end 전체 기여", lv=2)
tbl(doc,
    ["구분", "대상", "bag 개선", "핵심 수단"],
    [("PA01", "score_all 커널", "86,824→33,814 ms (−61%)", "CPU N4/OMP + CUDA hybrid T=256"),
     ("PA02", "matcher 오케스트레이션", "90,482→87,866 ms (−2.9%)", "make_cand + Branch + Score bucket"),],
    wids=[0.6, 1.45, 2.0, 2.5], sz=9)
P(doc,"",sb=2,sa=0)
P(doc,
  "PA02가 PA01보다 개선폭이 작은 이유: PA02 시작 시 이미 score_all(37.4%)이 최적화 완료되어 "
  "PA02 상한이 matcher_scope ~56.7 s로 제한됨. 그 안에서도 B&B 재귀·sort·호출 구조라는 "
  "알고리즘적 천장(Ceiling)이 존재하며 이것이 잔여 병목 53 s의 본질적 원인이다.",
  sz=10, sb=0, sa=3)

H(doc, "7.3  최종 프로덕션 빌드 플래그", lv=2)
for line in [
    "catkin_make -DPA01_OPT_LEVEL=9 -DPA01_USE_GPU=ON -DPA02_OPT_LEVEL=3  \\",
    "           -DPA01_GPU_THRESHOLD=256  \\",
    "           -DPA02_MAKE_CAND_OMP_MIN=512  \\",
    "           -DPA02_BRANCH_OMP_MIN=999999",
]:
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(0)
    cp.paragraph_format.left_indent  = Inches(0.25)
    ppr = cp._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),"F2F2F2")
    ppr.append(shd)
    r = cp.add_run(line)
    r.font.name="Consolas"; r.font.size=Pt(8.5)
    rpr2=r._element.get_or_add_rPr()
    rf2=OxmlElement("w:rFonts")
    rf2.set(qn("w:ascii"),"Consolas"); rf2.set(qn("w:hAnsi"),"Consolas")
    rpr2.append(rf2)
P(doc,"",sb=2,sa=0)

tbl(doc,
    ["파라미터", "값", "근거"],
    [("PA01_GPU_THRESHOLD", "256", "bag threshold sweep: T=256이 37.4 s로 최소 (T=64: 40.7 s, T=2048: 51.6 s)"),
     ("PA02_OPT_LEVEL", "3", "[match] L0→L3 −2.9%, upstream E1~E3 기각"),
     ("PA02_MAKE_CAND_OMP_MIN", "512", "hot path n=256에서 OMP 미발동, omp@256은 ×2.1 악화"),
     ("PA02_BRANCH_OMP_MIN", "999999", "GPU 빌드 concurrent CUDA unsafe — compile-time off"),],
    wids=[1.65, 0.55, 4.35], sz=8.5)
P(doc,"",sb=3,sa=0)

H(doc, "7.4  결론 및 향후 고속화 제언", lv=2)
P(doc,
  "본 과제는 단일 함수 최적화(PA01)에서 모듈 아키텍처 관점의 오케스트레이션 최적화(PA02)로 "
  "시각을 전환하여 matcher_scope 56,667→53,080 ms(−6.3%) 추가 단축을 달성하였다. "
  "특히 세 가지 독립적 실측 사례를 통해 '격리 마이크로벤치 승자 ≠ SLAM bag KPI 승자'임을 "
  "체계적으로 규명한 것은 임베디드 SLAM 시스템 성능 평가에 있어 방법론적 기여이다. "
  "Upstream ShrinkToFit의 semantics 불일치를 정확도 회귀 실측으로 잡아내 기각한 과정은 "
  "고성능 시스템에서 '빠르지만 틀린 결과'를 배제하는 품질 관리 원칙의 실증 사례이다.",
  sz=10, sb=0, sa=3)
P(doc,
  "잔여 병목 ~53 s는 B&B 재귀 트리 알고리즘 자체의 천장에 기인한다. "
  "추가 돌파구로는 ① Jetson LPDDR4 공유 메모리를 활용한 Zero-copy CUDA 포인터 아키텍처로 "
  "Host-Device 전송 오버헤드 근본 제거, ② 다중 CUDA 스트림 파이프라인 도입으로 "
  "복수 스캔 연산의 Waterfall 병렬화, ③ 멀티-해상도 grid를 device constant memory에 "
  "상주시켜 캐시 miss를 구조적으로 감소시키는 방향을 고려할 수 있다.",
  sz=10, sb=0, sa=4)

# ══════════════════════════════════════════
doc.save(str(OUT))
print("Saved:", OUT)
